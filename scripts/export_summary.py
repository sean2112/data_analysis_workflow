"""
data-analysis-workflow: 阶段总结导出工具（.xlsx + .docx）

每个分析阶段完成后调用此脚本，同时导出结构化数据 (.xlsx) 和可读报告 (.docx)。

格式分工:
  - .xlsx: 结构化/表格数据 (概览指标、列信息、对比数据、图表清单) — 适合进一步处理
  - .docx: 叙述性报告 (发现描述、分析结论、建议说明) — 适合阅读和分享
"""

import pandas as pd
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ═══════════════════════════════════════════════════════════════
#  .xlsx 导出函数
# ═══════════════════════════════════════════════════════════════

def export_phase_summary_xlsx(
    phase_number: int,
    phase_name: str,
    summary_data: dict,
    output_dir: str = ".",
) -> str:
    """导出阶段总结为 .xlsx 文件（结构化数据）。"""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_path = Path(output_dir) / f"{phase_number:01d}_{phase_name}_summary.xlsx"

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        if summary_data.get("overview"):
            pd.DataFrame(summary_data["overview"]).to_excel(
                writer, sheet_name="概览", index=False
            )
        if summary_data.get("details"):
            pd.DataFrame(summary_data["details"]).to_excel(
                writer, sheet_name="详细信息", index=False
            )
        if summary_data.get("findings"):
            pd.DataFrame({"发现": summary_data["findings"]}).to_excel(
                writer, sheet_name="发现与说明", index=False
            )
        # 元数据 sheet
        pd.DataFrame([
            {"字段": "阶段编号", "值": phase_number},
            {"字段": "阶段名称", "值": phase_name},
            {"字段": "导出时间", "值": datetime.now().strftime("%Y-%m-%d %H:%M:%S")},
        ]).to_excel(writer, sheet_name="元数据", index=False)

    return str(output_path)


def export_cleaning_report_xlsx(
    before: dict, after: dict, operations: list, output_dir: str = "."
) -> str:
    """导出数据清洗报告为 .xlsx。"""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_path = Path(output_dir) / "02_数据清洗_报告.xlsx"

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        comparison = pd.DataFrame([
            {"指标": k, "处理前": before.get(k, ""), "处理后": after.get(k, "")}
            for k in set(list(before.keys()) + list(after.keys()))
        ])
        comparison.to_excel(writer, sheet_name="清洗对比", index=False)
        pd.DataFrame({"操作步骤": operations}).to_excel(
            writer, sheet_name="操作记录", index=False
        )

    return str(output_path)


def export_final_report_xlsx(
    metrics: list[dict],
    findings: list[dict],
    recommendations: list[dict],
    output_dir: str = ".",
) -> str:
    """导出最终分析报告为 .xlsx。"""
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_path = Path(output_dir) / "06_最终报告.xlsx"

    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        if metrics:
            pd.DataFrame(metrics).to_excel(writer, sheet_name="关键指标", index=False)
        if findings:
            pd.DataFrame(findings).to_excel(writer, sheet_name="核心发现", index=False)
        if recommendations:
            pd.DataFrame(recommendations).to_excel(
                writer, sheet_name="建议方案", index=False
            )

    return str(output_path)


# ═══════════════════════════════════════════════════════════════
#  .docx 导出函数
# ═══════════════════════════════════════════════════════════════

def _add_styled_paragraph(doc, text, style="Normal", bold=False, size=None):
    """添加带样式的段落。"""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def _add_table_from_dicts(doc, data: list[dict], headers: list[str] = None):
    """从字典列表添加表格到文档。"""
    if not data:
        return
    if headers is None:
        headers = list(data[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)

    # 数据行
    for row_data in data:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            row_cells[i].text = str(row_data.get(h, ""))

    doc.add_paragraph()  # 表后空行


def export_phase_summary_docx(
    phase_number: int,
    phase_name: str,
    summary_data: dict,
    output_dir: str = ".",
) -> str | None:
    """导出阶段总结为 .docx 文件（叙述性报告）。"""
    if not HAS_DOCX:
        return None

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_path = Path(output_dir) / f"{phase_number:01d}_{phase_name}_总结报告.docx"
    doc = Document()

    # 标题
    title = doc.add_heading(f"阶段 {phase_number}：{phase_name}", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    # 概览表
    if summary_data.get("overview"):
        doc.add_heading("概览指标", level=2)
        _add_table_from_dicts(doc, summary_data["overview"])

    # 详细信息
    if summary_data.get("details"):
        doc.add_heading("详细信息", level=2)
        _add_table_from_dicts(doc, summary_data["details"])

    # 发现与说明
    if summary_data.get("findings"):
        doc.add_heading("发现与说明", level=2)
        for i, finding in enumerate(summary_data["findings"], 1):
            _add_styled_paragraph(doc, f"{i}. {finding}")

    # 页脚元数据
    doc.add_paragraph()
    doc.add_paragraph(f"阶段编号：{phase_number} | 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    doc.save(str(output_path))
    return str(output_path)


def export_cleaning_report_docx(
    before: dict, after: dict, operations: list, output_dir: str = "."
) -> str | None:
    """导出数据清洗报告为 .docx。"""
    if not HAS_DOCX:
        return None

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_path = Path(output_dir) / "02_数据清洗_总结报告.docx"
    doc = Document()

    doc.add_heading("数据清洗报告", level=1)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 清洗对比
    doc.add_heading("清洗前后对比", level=2)
    comparison = [
        {"指标": k, "处理前": before.get(k, ""), "处理后": after.get(k, "")}
        for k in set(list(before.keys()) + list(after.keys()))
    ]
    _add_table_from_dicts(doc, comparison)

    # 操作记录
    doc.add_heading("操作记录", level=2)
    for i, op in enumerate(operations, 1):
        _add_styled_paragraph(doc, f"{i}. {op}")

    doc.save(str(output_path))
    return str(output_path)


def export_final_report_docx(
    title: str,
    executive_summary: str,
    metrics: list[dict],
    findings: list[dict],
    recommendations: list[dict],
    output_dir: str = ".",
) -> str | None:
    """导出最终分析报告为 .docx（完整报告格式，适合分享）。"""
    if not HAS_DOCX:
        return None

    Path(output_dir).mkdir(parents=True, exist_ok=True)
    output_path = Path(output_dir) / "06_最终分析报告.docx"
    doc = Document()

    # 封面标题
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    # 执行摘要
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(executive_summary)

    # 关键指标
    if metrics:
        doc.add_heading("关键指标", level=1)
        _add_table_from_dicts(doc, metrics)

    # 核心发现
    if findings:
        doc.add_heading("核心发现", level=1)
        for f in findings:
            title_text = f.get("发现", f.get("title", ""))
            detail = f.get("说明", f.get("detail", ""))
            doc.add_heading(title_text, level=2)
            if detail:
                doc.add_paragraph(detail)

    # 建议方案
    if recommendations:
        doc.add_heading("建议方案", level=1)
        _add_table_from_dicts(doc, recommendations)

    doc.save(str(output_path))
    return str(output_path)


# ═══════════════════════════════════════════════════════════════
#  便捷组合导出
# ═══════════════════════════════════════════════════════════════

def export_phase_summary(
    phase_number: int,
    phase_name: str,
    summary_data: dict,
    output_dir: str = ".",
) -> dict:
    """
    同时导出 .xlsx + .docx 阶段总结。

    返回:
        {"xlsx": "路径", "docx": "路径或None"}
    """
    return {
        "xlsx": export_phase_summary_xlsx(phase_number, phase_name, summary_data, output_dir),
        "docx": export_phase_summary_docx(phase_number, phase_name, summary_data, output_dir),
    }


def export_cleaning_report(
    before: dict, after: dict, operations: list, output_dir: str = "."
) -> dict:
    """同时导出清洗报告的 .xlsx + .docx。"""
    return {
        "xlsx": export_cleaning_report_xlsx(before, after, operations, output_dir),
        "docx": export_cleaning_report_docx(before, after, operations, output_dir),
    }


def export_final_report(
    title: str,
    executive_summary: str,
    metrics: list[dict],
    findings: list[dict],
    recommendations: list[dict],
    output_dir: str = ".",
) -> dict:
    """同时导出最终报告的 .xlsx + .docx。"""
    return {
        "xlsx": export_final_report_xlsx(metrics, findings, recommendations, output_dir),
        "docx": export_final_report_docx(
            title, executive_summary, metrics, findings, recommendations, output_dir
        ),
    }
